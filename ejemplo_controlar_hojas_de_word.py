import docx
d = doc.Document(r'rutaejemplo')
d = docx.Document(r'rutaejemplo')
d.paragraphs
d.paragraphs[.text]
p = d.paragraphs[]
p.runs[.text]
p.runs[.bold.italic.underline]
p.runs[.underline = True]
d.save('rutaejempplo')
p.style
p.style = 'title'
p.save('rutaejemplo')



d = docx.Document()
d.add_paragraph('ejemplo')
d.add_paragraph('ejemplo')
d.save('rutaejemplo')
d.paragraphs[
p.add_run('ejemplo')
p.runs[.bold = True]
d.save('rutaejemplo')



def GetText(filename.doc):
    doc = docx.Document(filename.doc)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return 'slashn.'.join(fullText)
print(GetText('rutaejemplo'))

